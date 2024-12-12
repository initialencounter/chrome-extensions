from docx import Document
import json

def docx_to_text(docx_path):
    doc = Document(docx_path)
    full_text = set()  # 使用集合来避免重复
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:  # 忽略空段落
            full_text.add(text)
            
    return '\n'.join(full_text)

def docx_to_json(docx_path):
    doc = Document(docx_path)
    result = {
        'paragraphs': [],
        'tables': []
    }
    
    # 提取段落
    seen_paragraphs = set()
    for para in doc.paragraphs:
        text = para.text.strip()
        if text and text not in seen_paragraphs:
            result['paragraphs'].append({
                'text': text,
                'style': para.style.name
            })
            seen_paragraphs.add(text)
    
    # 提取表格
    for table in doc.tables:
        table_data = []
        seen_cells = set()
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text not in seen_cells:
                    row_data.append(cell_text)
                    seen_cells.add(cell_text)
            table_data.append(row_data)
        result['tables'].append(table_data)
    
    return json.dumps(result, ensure_ascii=False, indent=2)

# 使用示例
if __name__ == '__main__':
    file_path = ' 概要.docx'
    
    # 转换为JSON
    json_content = docx_to_json(file_path)
    print("JSON内容:", json_content)